# -*- coding: utf-8 -*-
"""把歷史新聞快照匯出成 Word（.docx）報告。

使用者在「📥 下載紀錄」頁自由選擇期間後，這裡負責把 data/news/ 底下
該區間的每日 JSON 快照組成一份可直接閱讀／存檔的 Word 文件。

只依賴 python-docx，不需要 Word 或 LibreOffice。
"""
import io
import json
import zipfile
from datetime import date, datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from . import store, topics as topics_mod

# Word 對中文要另外指定 eastAsia 字型，不然會用預設西文字型硬撐，字距很醜
ZH_FONT = "Microsoft JhengHei"
EN_FONT = "Calibri"

WEEKDAYS = ["一", "二", "三", "四", "五", "六", "日"]


# ── 日期與資料挑選 ──────────────────────────────────────────────
def parse_date(d):
    if isinstance(d, date):
        return d
    return datetime.strptime(str(d), "%Y-%m-%d").date()


def dates_in_range(start, end):
    """回傳區間內「實際有快照」的日期，由新到舊。"""
    s, e = parse_date(start), parse_date(end)
    if s > e:
        s, e = e, s
    return [d for d in store.available_dates() if s <= parse_date(d) <= e]


def weekday_label(date_str):
    return WEEKDAYS[parse_date(date_str).weekday()]


def _topic_meta():
    """topic_id → {name, emoji, category}，取自目前的 topics.json。"""
    return {t["id"]: {"name": t["name"],
                      "emoji": t.get("emoji", "📌"),
                      "category": topics_mod.category_of(t)}
            for t in store.load_topics()}


def topics_in_dates(dates):
    """這些日期的快照裡實際出現過的主題，回傳 [(topic_id, 顯示名稱), ...]。"""
    meta = _topic_meta()
    seen = {}
    for d in dates:
        snap = store.load_news(d) or {}
        names = snap.get("meta", {}).get("topic_names", {})
        for tid in snap.get("topics", {}):
            if tid not in seen:
                info = meta.get(tid)
                seen[tid] = info["name"] if info else names.get(tid, tid)
    # 依 topics.json 的順序排，已刪除的主題排在後面
    order = list(meta.keys())
    return sorted(seen.items(),
                  key=lambda kv: (order.index(kv[0]) if kv[0] in order else 999, kv[0]))


def collect(dates, topic_ids=None, only_official=False, limit_per_topic=0):
    """整理出 [(date_str, snapshot, [(topic_id, 名稱, brief, [article, ...]), ...]), ...]。"""
    meta = _topic_meta()
    out = []
    for d in dates:
        snap = store.load_news(d)
        if not snap:
            continue
        names = snap.get("meta", {}).get("topic_names", {})
        briefs = snap.get("meta", {}).get("briefs", {})
        blocks = []
        for tid, arts in snap.get("topics", {}).items():
            if topic_ids is not None and tid not in topic_ids:
                continue
            if only_official:
                arts = [a for a in arts if a.get("official")]
            if limit_per_topic:
                arts = arts[:limit_per_topic]
            info = meta.get(tid)
            name = info["name"] if info else names.get(tid, tid)
            emoji = info["emoji"] if info else "📌"
            blocks.append((tid, f'{emoji} {name}', briefs.get(tid, ""), arts))
        # 主題順序比照 topics.json
        order = list(meta.keys())
        blocks.sort(key=lambda b: order.index(b[0]) if b[0] in order else 999)
        out.append((d, snap, blocks))
    return out


# ── Word 排版工具 ──────────────────────────────────────────────
def _set_zh_font(run, size=None, bold=None, color=None):
    run.font.name = EN_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), ZH_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def _para(doc, text="", size=10.5, bold=False, color=None,
          space_after=4, indent=0, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Pt(indent)
    if text:
        _set_zh_font(p.add_run(text), size=size, bold=bold, color=color)
    return p


def _add_hyperlink(paragraph, text, url, size=11, bold=True):
    """python-docx 沒有超連結 API，得自己建立 relationship 與 w:hyperlink 節點。"""
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), EN_FONT)
    fonts.set(qn("w:hAnsi"), EN_FONT)
    fonts.set(qn("w:eastAsia"), ZH_FONT)
    rPr.append(fonts)
    if bold:
        rPr.append(OxmlElement("w:b"))
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    run.append(rPr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)
    return paragraph


def _init_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = EN_FONT
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), ZH_FONT)
    for name, size in (("Heading 1", 16), ("Heading 2", 13), ("Title", 22)):
        try:
            s = doc.styles[name]
        except KeyError:
            continue
        s.font.name = EN_FONT
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string("1F3864")
        s.element.rPr.rFonts.set(qn("w:eastAsia"), ZH_FONT)


def _overview_table(doc, data, topic_labels):
    """日期 × 主題的則數總表，讓人一眼看出哪幾天有料。"""
    cols = ["日期"] + [lbl for _, lbl in topic_labels] + ["小計"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, cols):
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        _set_zh_font(cell.paragraphs[0].add_run(text), size=9, bold=True)

    totals = {tid: 0 for tid, _ in topic_labels}
    for date_str, _snap, blocks in data:
        counts = {tid: len(arts) for tid, _n, _b, arts in blocks}
        row = table.add_row().cells
        _set_zh_font(row[0].paragraphs[0].add_run(date_str), size=9)
        for i, (tid, _lbl) in enumerate(topic_labels, start=1):
            n = counts.get(tid, 0)
            totals[tid] += n
            _set_zh_font(row[i].paragraphs[0].add_run(str(n)), size=9)
        _set_zh_font(row[-1].paragraphs[0].add_run(str(sum(counts.values()))), size=9, bold=True)

    row = table.add_row().cells
    _set_zh_font(row[0].paragraphs[0].add_run("合計"), size=9, bold=True)
    for i, (tid, _lbl) in enumerate(topic_labels, start=1):
        _set_zh_font(row[i].paragraphs[0].add_run(str(totals[tid])), size=9, bold=True)
    _set_zh_font(row[-1].paragraphs[0].add_run(str(sum(totals.values()))), size=9, bold=True)
    return table


def _render_article(doc, idx, a, include_summary=True, include_bullets=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(8)
    _set_zh_font(p.add_run(f"{idx}. "), size=11, bold=True)
    if a.get("url"):
        _add_hyperlink(p, a.get("title", "（無標題）"), a["url"])
    else:
        _set_zh_font(p.add_run(a.get("title", "（無標題）")), size=11, bold=True)

    badge = "官方來源" if a.get("official") else "媒體"
    meta = [badge, a.get("source_name", "")]
    if a.get("via") and a.get("via") != a.get("source_name"):
        meta.append(f'經由 {a["via"]}')
    if a.get("published_display"):
        meta.append(a["published_display"])
    _para(doc, "　".join(x for x in meta if x), size=9,
          color="6B7785", space_after=2, indent=14)

    if include_summary:
        text = (a.get("summary") or "").strip()
        if not text:
            text = "（此來源只提供標題，請點連結閱讀原文）"
        _para(doc, text, size=10, space_after=2, indent=14)

    if include_bullets:
        for b in a.get("bullets", []) or []:
            _para(doc, f"• {b}", size=9.5, color="3D4A57", space_after=1, indent=28)

    if a.get("url"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Pt(14)
        _set_zh_font(p.add_run("原始連結："), size=8.5, color="6B7785")
        _add_hyperlink(p, a["url"], a["url"], size=8.5, bold=False)


# ── 主要輸出 ────────────────────────────────────────────────────
def build_docx(dates, topic_ids=None, only_official=False, include_summary=True,
               include_bullets=True, include_overview=True, limit_per_topic=0,
               oldest_first=False):
    """把指定日期的快照組成一份 Word 報告，回傳 bytes。"""
    if not dates:
        raise ValueError("選取的期間內沒有任何快照。")

    ordered = sorted(dates, reverse=not oldest_first)
    data = collect(ordered, topic_ids=topic_ids, only_official=only_official,
                   limit_per_topic=limit_per_topic)
    if not data:
        raise ValueError("選取的期間內沒有任何快照。")

    total = sum(len(arts) for _d, _s, blocks in data for _tid, _n, _b, arts in blocks)
    official = sum(1 for _d, _s, blocks in data for _tid, _n, _b, arts in blocks
                   for a in arts if a.get("official"))
    span = f"{min(ordered)} ～ {max(ordered)}"

    doc = Document()
    _init_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_zh_font(title.add_run("📡 每日新聞雷達　新聞彙整"), size=22, bold=True)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_zh_font(sub.add_run(f"期間 {span}"), size=12, color="44546A")

    info = [f"共 {len(data)} 天", f"新聞 {total} 則", f"官方來源 {official} 則"]
    if only_official:
        info.append("（已篩選為僅官方來源）")
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_zh_font(meta_p.add_run("　｜　".join(info)), size=9.5, color="6B7785")
    gen = doc.add_paragraph()
    gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_zh_font(gen.add_run(f"匯出時間：{store.now_iso()}"), size=8.5, color="9AA4AF")

    if include_overview:
        doc.add_paragraph()
        _para(doc, "各日則數總表", size=12, bold=True, space_after=6)
        labels = topics_in_dates(ordered)
        if topic_ids is not None:
            labels = [(tid, lbl) for tid, lbl in labels if tid in topic_ids]
        if labels:
            _overview_table(doc, data, labels)

    for date_str, snap, blocks in data:
        doc.add_page_break()
        h = doc.add_heading(level=1)
        _set_zh_font(h.add_run(f"{date_str}（週{weekday_label(date_str)}）"),
                     size=16, bold=True)
        _para(doc, f'資料更新時間：{snap.get("generated_at", "—")}',
              size=8.5, color="9AA4AF", space_after=8)

        if not any(arts for _t, _n, _b, arts in blocks):
            _para(doc, "這一天沒有符合條件的新聞。", size=10, color="6B7785")
            continue

        for _tid, label, brief, arts in blocks:
            h2 = doc.add_heading(level=2)
            _set_zh_font(h2.add_run(f"{label}（{len(arts)} 則）"), size=13, bold=True)
            if brief:
                _para(doc, f"今日重點：{brief}", size=10, color="1F3864", space_after=6)
            if not arts:
                _para(doc, "（無符合條件的新聞）", size=9.5, color="6B7785")
                continue
            for i, a in enumerate(arts, start=1):
                _render_article(doc, i, a, include_summary=include_summary,
                                include_bullets=include_bullets)

    doc.add_page_break()
    _para(doc, "資料說明", size=12, bold=True, space_after=6)
    for line in (
        "本文件由「每日新聞雷達」自動彙整，內容來自各機關與媒體的公開 RSS 發布。",
        "摘要僅供快速瀏覽，正式引用請以原始連結的內容為準。",
        "標示「官方來源」者為政府機關、國際組織或發布方本身的網站。",
    ):
        _para(doc, f"• {line}", size=9.5, color="3D4A57", space_after=3)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_json_zip(dates):
    """把區間內的原始 JSON 快照打包成 zip，回傳 bytes。"""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        for d in sorted(dates, reverse=True):
            snap = store.load_news(d)
            if snap:
                z.writestr(f"news-{d}.json",
                           json.dumps(snap, ensure_ascii=False, indent=2))
    return buf.getvalue()


def filename(dates, ext="docx"):
    if not dates:
        return f"news-radar.{ext}"
    lo, hi = min(dates), max(dates)
    stem = lo if lo == hi else f"{lo}_to_{hi}"
    return f"news-radar_{stem}.{ext}"
